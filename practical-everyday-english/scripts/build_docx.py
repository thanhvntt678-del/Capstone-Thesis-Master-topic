#!/usr/bin/env python3
"""Build the PRACTICAL EVERYDAY ENGLISH manuscript (.docx) from lesson JSON.

Design (Rule 16 — learned from the SCM visual reference, layout only, never
its Vocabulary structure): navy/deep-blue hierarchy, pale-blue section bars,
a CEFR badge, clean A4 layout, a proper two-column-labelled dialogue table,
consistent footer, automatic page numbers. Long lessons flow onto a new
page — never shrunk or crammed.
"""
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
LESSONS_DIR = ROOT / "data" / "lessons"
OUT_PATH = ROOT / "manuscript" / "PRACTICAL_EVERYDAY_ENGLISH.docx"

NAVY = RGBColor(0x1F, 0x38, 0x64)
PALE_BLUE = "D9E2F3"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
BOOK_TITLE_EN = "PRACTICAL EVERYDAY ENGLISH"
BOOK_TITLE_VI = "TIẾNG ANH GIAO TIẾP THỰC CHIẾN"


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_footer(section, major_topic):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"PRACTICAL EVERYDAY ENGLISH  •  {major_topic}  •  page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_page_number_field(p)


def add_series_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(BOOK_TITLE_EN)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = NAVY
    run.font.name = "Arial"
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(BOOK_TITLE_VI)
    run2.font.size = Pt(10)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    run2.font.name = "Arial"
    p2.paragraph_format.space_after = Pt(12)


def add_cefr_badge_and_title(doc, lesson):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    badge_cell, title_cell = table.rows[0].cells
    badge_cell.width = Cm(3.2)
    title_cell.width = Cm(13.8)

    shade_cell(badge_cell, "1F3864")
    bp = badge_cell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run(lesson["cefr"])
    br.font.bold = True
    br.font.size = Pt(14)
    br.font.color.rgb = WHITE
    br.font.name = "Arial"
    bp2 = badge_cell.add_paragraph()
    bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br2 = bp2.add_run(f"LESSON {lesson['lesson_id']}")
    br2.font.size = Pt(8)
    br2.font.color.rgb = WHITE
    br2.font.name = "Arial"

    shade_cell(title_cell, PALE_BLUE)
    tp = title_cell.paragraphs[0]
    tr = tp.add_run(lesson["situation_title_en"])
    tr.font.bold = True
    tr.font.size = Pt(16)
    tr.font.color.rgb = NAVY
    tr.font.name = "Arial"
    tp2 = title_cell.add_paragraph()
    tr2 = tp2.add_run(lesson["situation_title_vi"])
    tr2.font.size = Pt(11)
    tr2.font.italic = True
    tr2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    tr2.font.name = "Arial"
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_section_bar(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  {text}")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    run.font.name = "Arial"
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), PALE_BLUE)
    pPr.append(shd)


def add_situation_context(doc, lesson):
    add_section_bar(doc, "SITUATION / BỐI CẢNH")
    p = doc.add_paragraph()
    r = p.add_run(lesson.get("context_note_en", ""))
    r.font.size = Pt(10.5)
    r.font.name = "Calibri"
    p2 = doc.add_paragraph()
    r2 = p2.add_run(lesson.get("context_note_vi", ""))
    r2.font.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    r2.font.name = "Calibri"


def add_dialogue_table(doc, lesson):
    add_section_bar(doc, "CORE DIALOGUE / HỘI THOẠI")
    turns = lesson["turns"]
    table = doc.add_table(rows=1 + len(turns), cols=4)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Cm(1.2), Cm(2.6), Cm(6.6), Cm(6.6)]
    headers = ["No.", "Speaker", "English", "Vietnamese Translation"]

    header_row = table.rows[0]
    for idx, (cell, header, width) in enumerate(zip(header_row.cells, headers, widths)):
        cell.width = width
        shade_cell(cell, "1F3864")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
        r.font.name = "Arial"

    for i, turn in enumerate(turns, start=1):
        row = table.rows[i]
        values = [str(turn["turn_order"]), turn["speaker"], turn["english"], turn["vietnamese"]]
        for cell, value, width in zip(row.cells, values, widths):
            cell.width = width
            p = cell.paragraphs[0]
            r = p.add_run(value)
            r.font.size = Pt(10.5)
            r.font.name = "Calibri"
        if turn["speaker"] == "Person B":
            for cell in row.cells:
                shade_cell(cell, "F2F5FB")


def build_lesson_section(doc, lesson, is_first):
    if not is_first:
        doc.add_section(WD_SECTION.NEW_PAGE)
        new_section = doc.sections[-1]
        new_section.left_margin = Cm(2.0)
        new_section.right_margin = Cm(2.0)
        new_section.top_margin = Cm(1.8)
        new_section.bottom_margin = Cm(1.8)
        set_footer(new_section, lesson["major_topic"])
    add_series_header(doc)
    add_cefr_badge_and_title(doc, lesson)
    add_situation_context(doc, lesson)
    add_dialogue_table(doc, lesson)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    lesson_files = sorted(LESSONS_DIR.glob("lesson_*.json"))
    lessons = [json.loads(p.read_text(encoding="utf-8")) for p in lesson_files]

    if lessons:
        set_footer(section, lessons[0]["major_topic"])

    for i, lesson in enumerate(lessons):
        build_lesson_section(doc, lesson, is_first=(i == 0))

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {len(lessons)} lessons to {OUT_PATH}")


if __name__ == "__main__":
    main()
