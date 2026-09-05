#!/usr/bin/env python3
"""Build the book manuscript under the MASTER V3 rules.

Inline bilingual layout: one paragraph per turn = bold character name +
English + natural Vietnamese translation, on the same visual row whenever
width permits (natural Word/LibreOffice line-wrap handles overflow — no
forced line breaks, no printed "English:"/"Vietnamese:" labels).

Book-only scope: this script produces the .docx manuscript. No Excel /
Python data export is generated at this stage (Master V3, Rule 1).
"""
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SITUATIONS_DIR = ROOT / "book" / "situations"
OUT_PATH = ROOT / "book" / "manuscript" / "PRACTICAL_EVERYDAY_ENGLISH_BOOK.docx"

NAVY = RGBColor(0x1F, 0x38, 0x64)
PALE_BLUE = "D9E2F3"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
VI_GREY = RGBColor(0x3A, 0x3A, 0x3A)
BOOK_TITLE_EN = "PRACTICAL EVERYDAY ENGLISH"
BOOK_TITLE_VI = "SÁCH GIAO TIẾP PHẢN XẠ THỰC TẾ"


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def set_footer(section, major_topic):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"PRACTICAL EVERYDAY ENGLISH  •  {major_topic}  •  page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_page_number_field(p)


def add_series_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(BOOK_TITLE_EN)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = NAVY
    r.font.name = "Arial"
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(BOOK_TITLE_VI)
    r2.font.size = Pt(9.5)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    r2.font.name = "Arial"
    p2.paragraph_format.space_after = Pt(8)


def add_situation_banner(doc, sit):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    badge_cell, title_cell = table.rows[0].cells
    badge_cell.width = Cm(3.2)
    title_cell.width = Cm(13.8)

    shade_cell(badge_cell, "1F3864")
    bp = badge_cell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run(sit["audited_cefr"])
    br.font.bold = True
    br.font.size = Pt(13)
    br.font.color.rgb = WHITE
    br.font.name = "Arial"
    bp2 = badge_cell.add_paragraph()
    bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br2 = bp2.add_run(f"SITUATION {sit['situation_id']}")
    br2.font.size = Pt(8)
    br2.font.color.rgb = WHITE
    br2.font.name = "Arial"

    shade_cell(title_cell, PALE_BLUE)
    tp = title_cell.paragraphs[0]
    tr = tp.add_run(sit["english_title"])
    tr.font.bold = True
    tr.font.size = Pt(15)
    tr.font.color.rgb = NAVY
    tr.font.name = "Arial"
    tp2 = title_cell.add_paragraph()
    tr2 = tp2.add_run(sit["vietnamese_title"])
    tr2.font.size = Pt(10.5)
    tr2.font.italic = True
    tr2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    tr2.font.name = "Arial"
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_context_note(doc, sit):
    en = sit.get("context_note_en", "")
    vi = sit.get("context_note_vi", "")
    if not en and not vi:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run(en)
    r1.font.italic = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r1.font.name = "Calibri"
    if vi:
        r2 = p.add_run("  " + vi)
        r2.font.italic = True
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        r2.font.name = "Calibri"


def add_dialogue(doc, sit):
    for turn in sit["turns"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.08

        name_run = p.add_run(f"{turn['speaker']}: ")
        name_run.font.bold = True
        name_run.font.color.rgb = NAVY
        name_run.font.size = Pt(10.5)
        name_run.font.name = "Calibri"

        en_run = p.add_run(turn["english"] + "  ")
        en_run.font.size = Pt(10.5)
        en_run.font.name = "Calibri"

        vi_run = p.add_run(turn["vietnamese"])
        vi_run.font.italic = True
        vi_run.font.size = Pt(10.5)
        vi_run.font.color.rgb = VI_GREY
        vi_run.font.name = "Calibri"


def build_situation_section(doc, sit, is_first):
    if not is_first:
        doc.add_section(WD_SECTION.NEW_PAGE)
        new_section = doc.sections[-1]
        new_section.left_margin = Cm(2.0)
        new_section.right_margin = Cm(2.0)
        new_section.top_margin = Cm(1.8)
        new_section.bottom_margin = Cm(1.8)
        set_footer(new_section, sit["major_topic"])
    add_series_header(doc)
    add_situation_banner(doc, sit)
    add_context_note(doc, sit)
    add_dialogue(doc, sit)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    situation_files = sorted(SITUATIONS_DIR.glob("situation_*.json"))
    situations = [json.loads(p.read_text(encoding="utf-8")) for p in situation_files]

    if situations:
        set_footer(section, situations[0]["major_topic"])

    for i, sit in enumerate(situations):
        build_situation_section(doc, sit, is_first=(i == 0))

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {len(situations)} situation(s) to {OUT_PATH}")


if __name__ == "__main__":
    main()
